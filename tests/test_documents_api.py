from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.documents.store import InMemoryDocumentStore
from app.main import create_app
from tests.fakes import BindableFakeListChatModel


def contract_bytes() -> bytes:
    document = Document()
    document.add_paragraph("示例合同")
    document.add_paragraph("第一条 付款")
    document.add_paragraph("甲方应当于十日内付款。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_upload_document_returns_summary() -> None:
    store = InMemoryDocumentStore()
    app = create_app(
        model=BindableFakeListChatModel(responses=["ok"]),
        document_store=store,
    )

    response = TestClient(app).post(
        "/documents",
        files={
            "file": (
                "示例合同.docx",
                contract_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201
    payload = response.json()
    assert payload["filename"] == "示例合同.docx"
    assert payload["section_count"] == 1
    assert store.get(payload["document_id"]).filename == "示例合同.docx"


def test_upload_rejects_non_docx() -> None:
    app = create_app(
        model=BindableFakeListChatModel(responses=["ok"]),
        document_store=InMemoryDocumentStore(),
    )

    response = TestClient(app).post(
        "/documents",
        files={"file": ("contract.txt", b"text", "text/plain")},
    )

    assert response.status_code == 415


def test_upload_rejects_empty_docx() -> None:
    app = create_app(
        model=BindableFakeListChatModel(responses=["ok"]),
        document_store=InMemoryDocumentStore(),
    )

    response = TestClient(app).post(
        "/documents",
        files={"file": ("empty.docx", b"", "application/octet-stream")},
    )

    assert response.status_code == 400


def test_upload_rejects_oversized_docx() -> None:
    app = create_app(
        model=BindableFakeListChatModel(responses=["ok"]),
        document_store=InMemoryDocumentStore(),
    )

    response = TestClient(app).post(
        "/documents",
        files={
            "file": (
                "large.docx",
                b"x" * (10 * 1024 * 1024 + 1),
                "application/octet-stream",
            )
        },
    )

    assert response.status_code == 413


def test_upload_rejects_broken_docx() -> None:
    app = create_app(
        model=BindableFakeListChatModel(responses=["ok"]),
        document_store=InMemoryDocumentStore(),
    )

    response = TestClient(app).post(
        "/documents",
        files={
            "file": (
                "broken.docx",
                b"not-a-word-file",
                "application/octet-stream",
            )
        },
    )

    assert response.status_code == 422