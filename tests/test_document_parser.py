from io import BytesIO

from docx import Document

from app.documents.parser import parse_docx


def build_contract_bytes() -> bytes:
    document = Document()
    document.add_paragraph("示例采购合同")
    document.add_paragraph("第一条 合同标的")
    document.add_paragraph("乙方向甲方提供测试设备。")
    document.add_paragraph("第二条 付款方式")
    document.add_paragraph("甲方应当于验收后十日内付款。")

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "合同金额"
    table.cell(0, 1).text = "1000元"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_parse_docx_preserves_sections_and_tables() -> None:
    parsed = parse_docx(build_contract_bytes(), "示例采购合同.docx")

    assert parsed.filename == "示例采购合同.docx"
    assert parsed.title == "示例采购合同"
    assert parsed.table_count == 1
    assert [(section.number, section.heading) for section in parsed.sections] == [
        ("第一条", "合同标的"),
        ("第二条", "付款方式"),
    ]
    assert "验收后十日内付款" in parsed.sections[1].content


def test_parse_docx_without_headings_falls_back_to_paragraph_sections() -> None:
    document = Document()
    document.add_paragraph("未使用标准标题的合同")
    document.add_paragraph("双方约定按月付款。")

    buffer = BytesIO()
    document.save(buffer)

    parsed = parse_docx(buffer.getvalue(), "非标准合同.docx")

    assert len(parsed.sections) == 2
    assert "未识别出显式条款标题" in parsed.warnings